import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    """Adds a heading with custom theme coloring and spacing."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    heading.paragraph_format.space_after = Pt(space_after)
    heading.paragraph_format.keep_with_next = True
    
    # Custom color settings for headings
    run = heading.runs[0]
    run.font.name = 'Outfit'
    if level == 1:
        run.font.color.rgb = RGBColor(124, 58, 237) # #7C3AED Deep Purple
        run.font.size = Pt(20)
    elif level == 2:
        run.font.color.rgb = RGBColor(139, 92, 246) # #8B5CF6 Purple
        run.font.size = Pt(14)
    else:
        run.font.color.rgb = RGBColor(17, 24, 39) # Dark Charcoal
        run.font.size = Pt(12)
    return heading

def main():
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("PROJECT PROPOSAL & TECHNICAL APPROACH")
    title_run.font.name = 'Outfit'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(124, 58, 237) # #7C3AED

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Secure Field Operations Mobile Platform (\"DeMars FieldOps Link\")")
    sub_run.font.name = 'Outfit'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(107, 114, 128) # Muted Gray

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.autofit = False
    
    meta_data = [
        ("Prepared For:", "Kyle Morris, President & Owner\nDeMars & Associates LTD"),
        ("Prepared By:", "GS3 Solutions Technical Architecture Team"),
        ("Project Name:", "DeMars FieldOps Link (Secure Transfer Agent Portal)"),
        ("Date:", "July 9, 2026")
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        
        cell_lbl = row.cells[0]
        cell_lbl.text = label
        cell_lbl.paragraphs[0].runs[0].font.bold = True
        cell_lbl.paragraphs[0].runs[0].font.name = 'Plus Jakarta Sans'
        cell_lbl.paragraphs[0].runs[0].font.size = Pt(10)
        cell_lbl.paragraphs[0].runs[0].font.color.rgb = RGBColor(107, 114, 128)
        cell_lbl.width = Inches(1.8)
        
        cell_val = row.cells[1]
        cell_val.text = val
        cell_val.paragraphs[0].runs[0].font.name = 'Plus Jakarta Sans'
        cell_val.paragraphs[0].runs[0].font.size = Pt(10)
        cell_val.paragraphs[0].runs[0].font.color.rgb = RGBColor(17, 24, 39)
        cell_val.width = Inches(4.7)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # 1. Executive Summary
    add_heading_styled(doc, "1. Executive Summary", level=1)
    
    p = doc.add_paragraph(
        "DeMars & Associates LTD is a nationally recognized alternative dispute resolution (ADR) firm. "
        "To execute warranty buyback agreements and Lemon Law dispute resolutions, DeMars deploys "
        "on-site Transfer Agents (vehicle retitling specialists) to retail dealerships or consumer locations. "
        "These agents are tasked with executing crucial steps: verifying vehicle identification and mileage, "
        "auditing cosmetic/structural conditions, obtaining owner signatures on retitling forms/vehicle titles, "
        "and exchanging legal manufacturer buyback checks."
    )
    p.paragraph_format.space_after = Pt(10)
    
    p2 = doc.add_paragraph(
        "GS3 Solutions proposes DeMars FieldOps Link, a custom-designed, highly secure mobile application "
        "specifically built for Transfer Agents in the field. The platform streamlines inspection logging, "
        "ensures title documents are validly executed, archives photographic proofs, and bridges real-time "
        "communication with the DeMars Home Office. Built with an offline-first architecture and military-grade "
        "encryption, DeMars FieldOps Link guarantees complete data integrity and compliance, even in areas with "
        "zero cell coverage."
    )
    p2.paragraph_format.space_after = Pt(18)

    # 2. Tech Stack & Platform Architecture
    add_heading_styled(doc, "2. Platform Architecture & Tech Stack", level=1)
    
    p = doc.add_paragraph(
        "To ensure portability, rapid updates, and enterprise security, the mobile app and cloud-backend "
        "are structured using modern, standardized frameworks:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    tech_table = doc.add_table(rows=6, cols=3)
    tech_table.style = 'Table Grid'
    
    headers = ["Architectural Layer", "Selected Technology Stack", "Business Rationale"]
    hdr_row = tech_table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.text = text
        set_cell_background(cell, "7C3AED") # Purple background
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Plus Jakarta Sans'
        run.font.size = Pt(10)

    tech_data = [
        ("Mobile Client", "React Native or Flutter", "Cross-platform consistency on iOS and Android. High performance with native device access (camera OCR, GPS, SQLCipher)."),
        ("Local Storage", "SQLite database via SQLCipher", "Encrypted, on-device data cache. Guarantees that sensitive consumer details and signatures remain secure at rest."),
        ("Cloud Backend", "Node.js (NestJS) or Go", "Extremely fast, type-safe REST APIs capable of handling real-time status updates and document processing queues."),
        ("Data Security", "AES-256 encryption at rest\nTLS 1.3 transit encryption", "Meets compliance standards for consumer financial/personal documentation. Protects titles and checks from intercept."),
        ("Cloud Hosting", "Amazon Web Services (AWS)", "Scaleable, SOC-2 compliant secure object buckets (S3) and RDS database instances with 99.99% operational uptime.")
    ]
    
    for row_idx, (layer, tech, rationale) in enumerate(tech_data, start=1):
        row = tech_table.rows[row_idx]
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = rationale
        
        # Style cells
        for cell in row.cells:
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Plus Jakarta Sans'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(17, 24, 39)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # 3. Core On-Site Buyback Workflows
    add_heading_styled(doc, "3. Core On-Site Buyback Workflows", level=1)
    
    p = doc.add_paragraph(
        "The application provides a structured, fail-safe wizard that enforces DeMars' audit guidelines "
        "for vehicle retitling and delivery transactions:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    workflows = [
        ("Step 1: Odometer & VIN Scan", "The agent uses the smartphone's camera to scan the vehicle's metal VIN plate. The app runs real-time OCR text recognition and automatically cross-references the VIN against the DeMars schedule. Odometer mileage is logged with a mandatory dashboard photo overlay."),
        ("Step 2: Condition & Key Checklist", "Agents complete a physical checklist verifying the cosmetic panel conditions, tires, and mechanical status. High-resolution photos are captured for each panel, embedding localized GPS and UTC timestamp metadata to prevent retroactive audit tampering."),
        ("Step 3: Document Verification", "The agent inspects and scans the customer's vehicle title (verifying signatures, release of bank liens, and power of attorney documents) and takes photographic proof of the check handed to the customer."),
        ("Step 4: Dual E-Signatures & Checksum Seal", "Once verification is completed, both the Transfer Agent and the customer sign the digital pads. The app captures coordinate points, signs the forms, and seals the record with a cryptographically secure SHA-256 hash.")
    ]
    
    for title, desc in workflows:
        wp = doc.add_paragraph()
        run_title = wp.add_run(f"•  {title}: ")
        run_title.font.bold = True
        run_title.font.name = 'Plus Jakarta Sans'
        run_title.font.color.rgb = RGBColor(139, 92, 246)
        
        run_desc = wp.add_run(desc)
        run_desc.font.name = 'Plus Jakarta Sans'
        run_desc.font.color.rgb = RGBColor(17, 24, 39)
        wp.paragraph_format.left_indent = Inches(0.25)
        wp.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # 4. Security & Cryptographic Integrity Controls
    add_heading_styled(doc, "4. Security & Cryptographic Integrity Controls", level=1)
    
    sec_items = [
        ("Offline-First SQL Sync", "Transfer Agents regularly work in dealership basements, steel-frame lots, or remote areas. The platform securely caches all forms, photographs, and signatures in a local SQL database. Once a web connection is recovered, background sync tasks push logs back to DeMars cloud storage."),
        ("SHA-256 Handover Checksums", "To satisfy manufacturer legal requirements, the application generates a local SHA-256 cryptographic hash of all combined data (VIN, odometer, photos, signatures) on submission. This checksum acts as a digital tamper-evident seal."),
        ("GPS Geofencing Validation", "The app records the device's precise GPS coordinate bounds at the time of e-signature, legally verifying that the transaction occurred at the correct dealership location."),
        ("Encrypted HQ Messenger", "A secure messaging gateway connecting agents directly to DeMars support. If a lien dispute or check discrepancy arises, agents can text support, upload documents, and get real-time clearance.")
    ]
    
    for title, desc in sec_items:
        sp = doc.add_paragraph()
        run_title = sp.add_run(f"•  {title}: ")
        run_title.font.bold = True
        run_title.font.name = 'Plus Jakarta Sans'
        run_title.font.color.rgb = RGBColor(124, 58, 237)
        
        run_desc = sp.add_run(desc)
        run_desc.font.name = 'Plus Jakarta Sans'
        run_desc.font.color.rgb = RGBColor(17, 24, 39)
        sp.paragraph_format.left_indent = Inches(0.25)
        sp.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # 5. Project Implementation Roadmap
    add_heading_styled(doc, "5. Project Implementation Roadmap", level=1)
    
    p = doc.add_paragraph(
        "GS3 Solutions proposes an agile 12-week development lifecycle to design, build, and deploy the application:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    roadmap_data = [
        ("Phase 1: UX/UI & Workflows", "Weeks 1 - 2", "Detailed wireframing of Transfer Agent screens. Alignment on Lemon Law and buyback legal requirements. Interactive mockups creation."),
        ("Phase 2: Mobile App Shell", "Weeks 3 - 5", "Setting up React Native / Flutter framework. Structuring local SQLite schemas with SQLCipher database encryption."),
        ("Phase 3: APIs & Sync Engines", "Weeks 6 - 8", "Building Node.js cloud APIs, AWS S3 image upload tunnels, and robust offline sync queue managers."),
        ("Phase 4: Integrations & OCR", "Weeks 9 - 10", "Integrating OCR camera parsing for VINs, building signature pads, and completing security penetration audits."),
        ("Phase 5: UAT & App Store Launch", "Weeks 11 - 12", "User acceptance testing with field agents. Finalizing compliance checks and submitting to Apple App Store and Google Play.")
    ]
    
    road_table = doc.add_table(rows=6, cols=3)
    road_table.style = 'Table Grid'
    
    hdr_row = road_table.rows[0]
    for idx, text in enumerate(["Development Phase", "Timeline", "Core Activities"]):
        cell = hdr_row.cells[idx]
        cell.text = text
        set_cell_background(cell, "7C3AED")
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Plus Jakarta Sans'
        run.font.size = Pt(10)
        
    for row_idx, (phase, timeline, desc) in enumerate(roadmap_data, start=1):
        row = road_table.rows[row_idx]
        row.cells[0].text = phase
        row.cells[1].text = timeline
        row.cells[2].text = desc
        
        # Style cells
        for cell in row.cells:
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Plus Jakarta Sans'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(17, 24, 39)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # 6. SLA & Maintenance
    add_heading_styled(doc, "6. Post-Launch Maintenance & Support SLA", level=1)
    
    p = doc.add_paragraph(
        "To ensure complete reliability in operational environments, GS3 Solutions provides continuous support:"
    )
    p.paragraph_format.space_after = Pt(10)
    
    sla_items = [
        ("Platform OS Upgrades", "Annual updates for iOS and Android, ensuring the application remains backward compatible and functional with new mobile releases."),
        ("Database Security Audits", "Biannual checks of SQLite and AWS server policies to prevent data leakage and maintain security standards."),
        ("24/7 Field Support Hotline", "A dedicated support line for Transfer Agents who encounter critical, blocker errors while executing vehicle buybacks on-site.")
    ]
    
    for title, desc in sla_items:
        ap = doc.add_paragraph()
        run_title = ap.add_run(f"•  {title}: ")
        run_title.font.bold = True
        run_title.font.name = 'Plus Jakarta Sans'
        run_title.font.color.rgb = RGBColor(107, 114, 128)
        
        run_desc = ap.add_run(desc)
        run_desc.font.name = 'Plus Jakarta Sans'
        run_desc.font.color.rgb = RGBColor(17, 24, 39)
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(6)

    # Save Document
    filename = "DeMars_FieldOps_App_Proposal.docx"
    doc.save(filename)
    print(f"Proposal document saved as '{filename}' successfully!")

if __name__ == "__main__":
    main()
